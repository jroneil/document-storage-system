from PyPDF2 import PdfReader
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from keybert import KeyBERT
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from textblob import TextBlob
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
import numpy as np

kw_model = KeyBERT()

def extract_text_from_pdf(file_path):
    # Try extracting text directly
    reader = PdfReader(file_path)
    text = ''
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:  # If text is extracted successfully
            text += page_text
        else:  # Fallback to OCR for scanned PDFs
            images = convert_from_path(file_path)
            for image in images:
                text += pytesseract.image_to_string(image)
    return text

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_keywords(text):
    keywords = kw_model.extract_keywords(text, keyphrase_ngram_range=(1, 2), stop_words='english')
    return [kw[0] for kw in keywords]

def generate_summary(text, sentences_count=3):
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = LsaSummarizer()
    summary = summarizer(parser.document, sentences_count)
    return " ".join([str(sentence) for sentence in summary])

def analyze_sentiment(text):
    blob = TextBlob(text)
    return str(blob.sentiment)

def generate_tags(keywords, num_tags=3):
    vectorizer = TfidfVectorizer()
    X = vectorizer.fit_transform(keywords)
    kmeans = KMeans(n_clusters=num_tags)
    kmeans.fit(X)
    tags = []
    for i in range(num_tags):
        centroid = np.argsort(kmeans.cluster_centers_[i])[-1]
        tags.append(keywords[centroid])
    return tags

from .rabbitmq import publish_message

def process_file(file_path):
    # Extract text based on file type
    if file_path.endswith('.pdf'):
        text = extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        text = extract_text_from_docx(file_path)
    else:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()

    # Extract keywords
    keywords = extract_keywords(text)

    # Generate tags
    tags = generate_tags(keywords)

    # Generate summary
    summary = generate_summary(text)

    # Analyze sentiment
    sentiment = analyze_sentiment(text)

    # Save metadata to the database
    document = Document(
        filename=os.path.basename(file_path),
        tags=",".join(tags),
        keywords=",".join(keywords),
        summary=summary,
        sentiment=sentiment
    )
    db.session.add(document)
    db.session.commit()

    # Publish a message to RabbitMQ
    message = {
        'filename': document.filename,
        'tags': document.tags,
        'keywords': document.keywords,
        'summary': document.summary,
        'sentiment': document.sentiment
    }
    publish_message(message)

    # Move the file to the uploads folder
    os.rename(file_path, os.path.join(UPLOAD_FOLDER, os.path.basename(file_path)))